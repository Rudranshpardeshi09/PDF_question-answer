"""
Syllabus Service - Parses PDF/DOCX syllabi with table support
Extracts: Subject Code (IC-XXX format), Unit No, Title, Contents
"""

import tempfile
import re
from typing import List, Dict, Optional, Tuple
from pypdf import PdfReader
from docx import Document


# ══════════════════════════════════════════════════════════════════════════════
# TEXT EXTRACTION
# ══════════════════════════════════════════════════════════════════════════════

def extract_text_from_pdf(path: str) -> Tuple[str, List[List[List[str]]]]:
    """
    Extract text and table-like structures from PDF.
    Returns: (plain_text, tables)
    """
    reader = PdfReader(path)
    text_parts = []
    tables = []
    
    for page in reader.pages:
        page_text = page.extract_text() or ""
        text_parts.append(page_text)
        
        # Try to detect table-like structures from text
        lines = page_text.split('\n')
        current_table = []
        
        for line in lines:
            # Detect table rows (pipe-separated or tab-separated)
            if '|' in line or '\t' in line:
                cells = re.split(r'\s*\|\s*|\t+', line.strip())
                cells = [c.strip() for c in cells if c.strip()]
                if len(cells) >= 2:
                    current_table.append(cells)
            elif current_table:
                if len(current_table) >= 2:
                    tables.append(current_table)
                current_table = []
        
        if current_table and len(current_table) >= 2:
            tables.append(current_table)
    
    return '\n'.join(text_parts), tables


def extract_tables_from_docx(path: str) -> Tuple[str, List[List[List[str]]]]:
    """
    Extract text and tables from DOCX with proper table parsing.
    Returns: (plain_text, tables)
    """
    doc = Document(path)
    text_parts = []
    tables = []
    
    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())
    
    # Extract tables properly
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_cells.append(cell_text)
            if any(c for c in row_cells):  # Skip empty rows
                table_data.append(row_cells)
        
        if len(table_data) >= 2:  # At least header + 1 row
            tables.append(table_data)
    
    return '\n'.join(text_parts), tables


def parse_syllabus(file):
    """
    Parse syllabus from PDF or DOCX and extract structured data:
    - Subject code (IC-XXX format like IC-812 Theory of Computation)
    - Units with numbers and titles from table
    - Contents/Topics (comma-separated in cells)
    
    Returns:
    {
        "subject": "IC-812 Theory of Computation",
        "units": [
            {
                "name": "Unit 1: Introduction",
                "title": "Introduction",
                "topics": ["topic1", "topic2", ...],
                "format": "long"
            }
        ]
    }
    """
    import os
    
    path = None
    
    try:
        # Save uploaded file safely
        with tempfile.NamedTemporaryFile(delete=False, suffix='.tmp') as tmp:
            tmp.write(file.file.read())
            path = tmp.name

        filename = file.filename.lower()
        text = ""
        tables = []

        if filename.endswith(".pdf"):
            text, tables = extract_text_from_pdf(path)
        elif filename.endswith(".docx"):
            text, tables = extract_tables_from_docx(path)
        else:
            raise ValueError("Unsupported syllabus format. Please upload PDF or DOCX.")

        # ─────── EXTRACT SUBJECT CODE (IC-XXX format) ───────
        subject = extract_subject_code(text)

        # ─────── EXTRACT UNITS & TOPICS FROM TABLES ───────
        units = extract_units_from_tables(tables, text)

        # Fallback to text-based extraction if no table units found
        if not units:
            units = extract_units_from_text(text)

        # Final fallback
        if not units:
            lines = [line.strip() for line in text.splitlines() if len(line.strip()) > 10]
            units = [{
                "name": "General Topics",
                "title": "General Topics",
                "topics": lines[:50],
                "format": "long"
            }]

        return {
            "subject": subject,
            "units": units
        }
    
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(
            "Failed to read syllabus file. Please upload a valid PDF or DOCX."
        ) from e
    finally:
        # ✅ CLEANUP: Remove temp file to free disk space
        if path and os.path.exists(path):
            try:
                os.remove(path)
            except OSError:
                pass  # Ignore cleanup errors


# ══════════════════════════════════════════════════════════════════════════════
# SUBJECT CODE EXTRACTION
# ══════════════════════════════════════════════════════════════════════════════

def extract_subject_code(text: str) -> str:
    """
    Extract subject code and name from syllabus.
    Targets format: IC-812 Theory of Computation (IC- is permanent, rest varies)
    Also handles: CS-XXX, IT-XXX, EC-XXX, ME-XXX etc.
    """
    # Pattern for subject codes like IC-812 Theory of Computation
    subject_patterns = [
        # IC-812 Theory of Computation format
        r'([A-Z]{2,4}[-\s]?\d{3,4}[A-Z]?)\s*[:\-–]?\s*([A-Za-z][A-Za-z\s&,\-]+)',
        # Subject Code: IC-812
        r'(?:Subject\s*Code|Course\s*Code|Code)\s*[:\-–]?\s*([A-Z]{2,4}[-\s]?\d{3,4}[A-Z]?)',
        # Subject Name after code pattern
        r'(?:Subject\s*Name|Course\s*Name|Course\s*Title)\s*[:\-–]?\s*([A-Za-z][A-Za-z\s&,\-]+)',
    ]
    
    subject_code = ""
    subject_name = ""
    
    # Try to extract subject code with name (like "IC-812 Theory of Computation")
    pattern1 = r'([A-Z]{2,4}[-\s]?\d{3,4}[A-Z]?)\s*[:\-–]?\s*([A-Za-z][A-Za-z\s&,\-]{3,50})'
    match = re.search(pattern1, text)
    if match:
        subject_code = match.group(1).strip()
        subject_name = match.group(2).strip()
        # Clean up the name
        subject_name = re.sub(r'\s+', ' ', subject_name).strip()
        return f"{subject_code} {subject_name}"[:100]
    
    # Try individual patterns
    for pattern in subject_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            result = match.group(1).strip()
            if match.lastindex >= 2:
                result = f"{result} {match.group(2).strip()}"
            return result[:100]
    
    # Fallback: first meaningful line
    for line in text.split('\n')[:10]:
        line = line.strip()
        if len(line) > 5 and not line.lower().startswith(('unit', 'module', 'chapter')):
            return line[:100]
    
    return "Unknown Subject"


# ══════════════════════════════════════════════════════════════════════════════
# TABLE-BASED EXTRACTION
# ══════════════════════════════════════════════════════════════════════════════

def find_column_indices(header_row: List[str]) -> Dict[str, int]:
    """
    Find column indices for Unit No, Title, Contents from header row.
    Handles various column name variations.
    """
    indices = {"unit_no": -1, "title": -1, "contents": -1}
    
    # Column name variations
    unit_names = ['unit no', 'unit', 'unit no.', 'module', 'module no', 'sl', 'sno', 's.no', 'no', 'sr']
    title_names = ['title', 'topic', 'unit title', 'module title', 'name', 'chapter', 'heading']
    content_names = ['contents', 'content', 'topics', 'syllabus', 'description', 'details', 'sub-topics', 'subtopics']
    
    for idx, cell in enumerate(header_row):
        cell_lower = cell.lower().strip()
        
        if any(name in cell_lower for name in unit_names) and indices["unit_no"] == -1:
            indices["unit_no"] = idx
        elif any(name in cell_lower for name in title_names) and indices["title"] == -1:
            indices["title"] = idx
        elif any(name in cell_lower for name in content_names) and indices["contents"] == -1:
            indices["contents"] = idx
    
    return indices


def parse_contents_cell(content: str) -> List[str]:
    """
    Parse comma-separated contents from a table cell.
    Handles formats like: "Wireless Networks, Wireless vs Wired Networks, mobile devices"
    """
    if not content:
        return []
    
    # Split by comma, semicolon, or newline
    topics = re.split(r'[,;]\s*|\n+', content)
    
    # Clean and filter
    cleaned_topics = []
    for topic in topics:
        topic = topic.strip()
        # Remove numbering like "1.", "a)", "(i)" etc.
        topic = re.sub(r'^[\d]+[.\)]\s*|^[a-z][.\)]\s*|^\([ivxlc]+\)\s*', '', topic, flags=re.IGNORECASE)
        topic = topic.strip()
        
        if len(topic) >= 3:  # Minimum meaningful length
            cleaned_topics.append(topic)
    
    return cleaned_topics


def extract_units_from_tables(tables: List[List[List[str]]], text: str) -> List[Dict]:
    """
    Extract units and topics from detected tables.
    Looks for tables with columns: Unit No, Title, Contents
    """
    units = []
    
    for table in tables:
        if len(table) < 2:  # Need header + at least 1 data row
            continue
        
        # Try to find header row (first row with column-like names)
        header_idx = 0
        indices = find_column_indices(table[0])
        
        # If first row doesn't look like header, try second row
        if indices["contents"] == -1 and len(table) > 1:
            indices = find_column_indices(table[1])
            if indices["contents"] != -1:
                header_idx = 1
        
        # Skip if no content column found
        if indices["contents"] == -1:
            continue
        
        # Process data rows
        for row_idx, row in enumerate(table[header_idx + 1:], start=1):
            if len(row) <= max(i for i in indices.values() if i >= 0):
                continue
            
            # Extract unit number
            unit_no = ""
            if indices["unit_no"] >= 0 and indices["unit_no"] < len(row):
                unit_no = row[indices["unit_no"]].strip()
                # Normalize unit number (I, II, III or 1, 2, 3)
                unit_no = re.sub(r'^unit\s*', '', unit_no, flags=re.IGNORECASE).strip()
            
            if not unit_no:
                unit_no = str(row_idx)
            
            # Extract title
            title = ""
            if indices["title"] >= 0 and indices["title"] < len(row):
                title = row[indices["title"]].strip()
            
            # Extract contents/topics
            contents = ""
            if indices["contents"] >= 0 and indices["contents"] < len(row):
                contents = row[indices["contents"]]
            
            topics = parse_contents_cell(contents)
            
            if topics:  # Only add if we have topics
                unit_name = f"Unit {unit_no}"
                if title:
                    unit_name = f"Unit {unit_no}: {title}"
                
                units.append({
                    "name": unit_name,
                    "title": title or f"Unit {unit_no}",
                    "topics": topics,
                    "format": infer_format(topics)
                })
    
    return units


def extract_units_from_text(text: str) -> List[Dict]:
    """
    Fallback: Extract units and topics from plain text when no table is found.
    """
    units = []
    current_unit = None
    current_title = ""
    current_topics = []

    lines = text.split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Detect Unit header (Unit I, Unit II, Unit 1, Unit 2, Module 1, etc.)
        unit_match = re.match(
            r'(?:unit|module|chapter)\s*([ivxlc0-9]+)\s*[:\-–]?\s*(.*)',
            line, re.IGNORECASE
        )

        if unit_match:
            # Save previous unit
            if current_unit and current_topics:
                units.append({
                    "name": current_unit,
                    "title": current_title,
                    "topics": current_topics,
                    "format": infer_format(current_topics)
                })

            # Start new unit
            unit_num = unit_match.group(1).upper()
            current_title = unit_match.group(2).strip() if unit_match.group(2) else ""
            current_unit = f"Unit {unit_num}"
            if current_title:
                current_unit = f"Unit {unit_num}: {current_title}"
            current_topics = []

        # Add topics to current unit (comma-separated or individual lines)
        elif current_unit:
            # Check if line contains comma-separated topics
            if ',' in line and len(line) > 10:
                topics = parse_contents_cell(line)
                current_topics.extend(topics)
            elif len(line) > 4 and not re.match(r'^[\d\-•*]+$', line):
                current_topics.append(line)

    # Don't forget last unit
    if current_unit and current_topics:
        units.append({
            "name": current_unit,
            "title": current_title,
            "topics": current_topics,
            "format": infer_format(current_topics)
        })

    return units


def infer_format(topics: List[str]) -> str:
    """
    Infer answer format based on topic complexity.
    """
    if not topics:
        return "short"

    avg_length = sum(len(t) for t in topics) / len(topics)

    if avg_length > 100:
        return "long"
    elif avg_length > 40:
        return "medium"
    else:
        return "short"
