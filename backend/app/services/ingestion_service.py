# this file handles loading and processing PDFs into searchable chunks
import logging
import os
import shutil

from fastapi import UploadFile
from langchain_community.document_loaders import PyPDFLoader

from app.rag.chunking import chunk_documents
from app.vectorstore.faiss_store import save_vectorstore

logger = logging.getLogger(__name__)

# folder where all uploaded PDFs are stored
UPLOAD_DIR = "app/data/uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)


def _load_documents(persistent_path: str):
    if persistent_path.lower().endswith(".pdf"):
        loader = PyPDFLoader(persistent_path)
        if hasattr(loader, "lazy_load"):
            return list(loader.lazy_load())
        return loader.load()

    if persistent_path.lower().endswith(".docx"):
        from docx import Document as DocxDocument
        from langchain_core.documents import Document

        doc = DocxDocument(persistent_path)
        full_text = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        return [Document(page_content="\n".join(full_text), metadata={"source": persistent_path})]

    raise ValueError("Unsupported file format")


# takes a PDF file and processes it into searchable chunks stored in our database
def ingest_pdf(input_source):
    if isinstance(input_source, UploadFile):
        filename = input_source.filename
        persistent_path = os.path.join(UPLOAD_DIR, filename)
        with open(persistent_path, "wb") as f:
            shutil.copyfileobj(input_source.file, f)
    elif isinstance(input_source, str):
        persistent_path = input_source
        filename = os.path.basename(persistent_path)
        if not os.path.exists(persistent_path):
            raise FileNotFoundError(f"PDF not found: {persistent_path}")
    else:
        raise TypeError("ingest_pdf expects UploadFile or file path")

    try:
        logger.info("Loading PDF: %s", filename)
        documents = _load_documents(persistent_path)

        if not documents:
            raise ValueError(f"PDF file '{filename}' is empty or unreadable - no text could be extracted")

        total_pages = len(documents)
        logger.info("Loaded %s pages from %s", total_pages, filename)

        chunks = chunk_documents(documents)
        if not chunks:
            raise ValueError("No chunks created from document")
        logger.info("Created %s chunks", len(chunks))

        save_vectorstore(chunks)
        logger.info("Successfully ingested %s", filename)

        return {
            "status": "success",
            "filename": filename,
            "pages": total_pages,
            "chunks": len(chunks),
        }

    except Exception as e:
        logger.error("Error processing PDF %s: %s", filename, e)
        raise


# rebuilds the entire vector database from all remaining PDFs
# this is called after deleting a PDF to keep the database accurate
def rebuild_vectorstore_from_uploads():
    from app.vectorstore.faiss_store import replace_vectorstore
    from app.core.config import settings

    uploads_dir = "app/data/uploads"

    if not os.path.exists(uploads_dir):
        return

    pdf_files = [
        f for f in os.listdir(uploads_dir)
        if f.lower().endswith(".pdf") or f.lower().endswith(".docx")
    ]

    if not pdf_files:
        if os.path.exists(settings.VECTOR_DB_PATH):
            shutil.rmtree(settings.VECTOR_DB_PATH)
        logger.info("No PDFs remaining, vectorstore cleared")
        return

    documents = []

    for filename in pdf_files:
        try:
            path = os.path.join(uploads_dir, filename)
            documents.extend(_load_documents(path))
            logger.info("Loaded %s for rebuild", filename)
        except Exception as e:
            logger.warning("Failed to load %s: %s", filename, e)

    if not documents:
        if os.path.exists(settings.VECTOR_DB_PATH):
            shutil.rmtree(settings.VECTOR_DB_PATH)
        return

    chunks = chunk_documents(documents)
    replace_vectorstore(chunks)
    logger.info("Rebuilt vectorstore with %s chunks from %s PDFs", len(chunks), len(pdf_files))
